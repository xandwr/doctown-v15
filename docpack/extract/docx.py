"""DOCX document extraction using python-docx."""

from __future__ import annotations

import io
import zipfile
from xml.etree import ElementTree as ET

from .base import ExtractedDocument, ExtractedImage


def extract_docx(data: bytes) -> ExtractedDocument:
    """
    Extract text and images from DOCX bytes.

    Args:
        data: Raw DOCX file bytes

    Returns:
        ExtractedDocument with text content and embedded images
    """
    from docx import Document

    doc = Document(io.BytesIO(data))
    text_parts: list[str] = []
    images: list[ExtractedImage] = []
    current_heading: str | None = None

    # Extract text from paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.append(text)

            # Track headings for context
            if para.style and para.style.name and "Heading" in para.style.name:
                current_heading = text

    # Extract text from tables
    for table in doc.tables:
        table_rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                table_rows.append(" | ".join(cells))
        if table_rows:
            text_parts.append("\n".join(table_rows))

    # Extract images from the document
    image_map = _extract_images_from_docx(data)
    image_index = 0

    for rel_id, (img_data, content_type) in image_map.items():
        # Determine format from content type
        if "png" in content_type:
            fmt = "png"
        elif "jpeg" in content_type or "jpg" in content_type:
            fmt = "jpeg"
        elif "gif" in content_type:
            fmt = "gif"
        else:
            fmt = "unknown"

        # Get dimensions using PIL
        width, height = 0, 0
        try:
            from PIL import Image

            pil_img = Image.open(io.BytesIO(img_data))
            width, height = pil_img.size
        except Exception:
            pass

        images.append(
            ExtractedImage(
                data=img_data,
                format=fmt,
                width=width,
                height=height,
                page_number=None,  # DOCX doesn't have real pages
                image_index=image_index,
                context=current_heading,
            )
        )
        image_index += 1

    # Get metadata
    metadata: dict[str, str] = {}
    try:
        if doc.core_properties:
            if doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
            if doc.core_properties.author:
                metadata["author"] = doc.core_properties.author
    except Exception:
        pass

    return ExtractedDocument(
        text="\n\n".join(text_parts),
        images=images,
        page_count=None,
        metadata=metadata,
    )


def _extract_images_from_docx(data: bytes) -> dict[str, tuple[bytes, str]]:
    """Extract all images from DOCX archive."""
    images: dict[str, tuple[bytes, str]] = {}

    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            # Parse relationships to find image files
            try:
                rels_xml = zf.read("word/_rels/document.xml.rels")
                root = ET.fromstring(rels_xml)

                for rel in root:
                    target = rel.get("Target", "")
                    rel_id = rel.get("Id", "")

                    if "media/" in target:
                        try:
                            # Handle relative paths
                            if target.startswith("/"):
                                img_path = target[1:]
                            else:
                                img_path = f"word/{target}"

                            img_data = zf.read(img_path)

                            # Determine content type from extension
                            if target.endswith(".png"):
                                content_type = "image/png"
                            elif target.endswith((".jpg", ".jpeg")):
                                content_type = "image/jpeg"
                            elif target.endswith(".gif"):
                                content_type = "image/gif"
                            else:
                                content_type = "image/unknown"

                            images[rel_id] = (img_data, content_type)
                        except KeyError:
                            continue
            except KeyError:
                pass
    except zipfile.BadZipFile:
        pass

    return images
